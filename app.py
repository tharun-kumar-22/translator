import streamlit as st
from docx import Document
from googletrans import Translator
from io import BytesIO
import zipfile
from datetime import datetime

st.title("Word Document Translator powered by Pfeuffer GmbH")
st.write("Upload multilingual Word documents, edit the German version, and download all updated versions.")

# Step 1: Upload Documents
uploaded_files = st.file_uploader(
    "Upload Word Documents (One per language)",
    type=["docx"],
    accept_multiple_files=True
)

# Step 2: Allow Manual Language Selection
language_files = {}
if uploaded_files:
    st.write("Please select the language for each uploaded file:")
    for file in uploaded_files:
        language = st.selectbox(f"Select language for: {file.name}", 
                                options=["German", "English", "French", "Spanish"], 
                                key=file.name)
        language_files[language.lower()] = file
    st.success(f"Uploaded and mapped {len(language_files)} files.")

# Step 3: Extract German Text for Editing
def extract_text_from_docx(file):
    """Extract text from a DOCX file."""
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

german_text = None
if "german" in language_files:
    german_file = language_files["german"]
    german_text = extract_text_from_docx(german_file)
    edited_german_text = st.text_area("Edit German Document", german_text, key="edit_german")
else:
    st.error("Please assign a file as 'German' for editing.")

# Step 4: Translate Edits
translator = Translator()
translated_documents = {}

if german_text and edited_german_text:
    st.write("Translated Documents Preview:")

    for lang_code, lang_name in [("en", "English"), ("fr", "French"), ("es", "Spanish")]:  # Add more languages if needed
        translated_text = translator.translate(edited_german_text, src='de', dest=lang_code).text
        st.text_area(f"{lang_name} Version", translated_text, disabled=True)
        translated_documents[lang_code] = translated_text

# Step 5: Save and Download Updated Documents
def create_docx(text, file_name):
    """Create a DOCX file from text and return as a byte stream."""
    doc = Document()
    doc.add_paragraph(text)
    byte_stream = BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)
    return byte_stream

if st.button("Save All and Download"):
    download_links = {}

    # Save the edited German document
    if edited_german_text:
        download_links["german_updated.docx"] = create_docx(edited_german_text, "german_updated.docx")

    # Save the translated documents
    for lang_code, lang_text in translated_documents.items():
        file_name = f"{lang_code}_updated.docx"
        download_links[file_name] = create_docx(lang_text, file_name)

    # Bundle all updated files into a ZIP
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zip_file:
        for file_name, file_stream in download_links.items():
            zip_file.writestr(file_name, file_stream.getvalue())
    zip_buffer.seek(0)

    st.download_button(
        label="Download All Documents (ZIP)",
        data=zip_buffer,
        file_name=f"Updated_Documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
        mime="application/zip"
    )
